#!/usr/bin/env python3
"""Gera templates DOCX para procuração e contrato de honorários."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


TEMPLATES_DIR = Path(__file__).parent


def _set_margins(doc: Document, top: float, bottom: float, left: float, right: float) -> None:
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _add_paragraph(doc: Document, text: str, bold: bool = False, align: int | None = None, size: int = 12, spacing_after: float = 0.5) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after * 12)
    p.paragraph_format.line_spacing = 1.5


def create_procuracao() -> None:
    doc = Document()
    _set_margins(doc, top=3, bottom=2, left=3, right=2)

    # Título
    _add_paragraph(doc, "PROCURAÇÃO AD JUDICIA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, spacing_after=2)

    # Corpo
    body = (
        "Pelo presente instrumento particular de procuração, "
        "{{ cliente_nome }}, {{ cliente_nacionalidade }}, {{ cliente_estado_civil }}, "
        "{{ cliente_profissao }}, portador(a) do RG nº {{ cliente_rg }} e "
        "inscrito(a) no CPF/CNPJ sob nº {{ cliente_cpf }}, "
        "residente e domiciliado(a) em {{ cliente_endereco }}, "
        "nomeia e constitui seu bastante procurador(a) o(a) Dr(a). "
        "{{ advogado_nome }}, inscrito(a) na OAB sob nº {{ advogado_oab }}, "
        "com escritório profissional na {{ advogado_endereco }}, "
        "a quem confere amplos poderes para o foro em geral, conforme Art. 105 "
        "do Código de Processo Civil, podendo propor ações, contestar, "
        "recorrer, transigir, desistir, acordar, receber e dar quitação, "
        "firmar compromisso, substabelecer com ou sem reserva de poderes, "
        "e praticar todos os demais atos necessários ao bom e fiel "
        "desempenho deste mandato, {{ poderes }}."
    )
    _add_paragraph(doc, body)

    _add_paragraph(doc, "", spacing_after=2)  # espaço

    # Local e data
    _add_paragraph(doc, "{{ cidade }}, {{ data_extenso }}.", align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_paragraph(doc, "", spacing_after=3)  # espaço para assinatura

    # Assinatura
    _add_paragraph(doc, "_" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
    _add_paragraph(doc, "{{ cliente_nome }}", align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
    _add_paragraph(doc, "Outorgante", align=WD_ALIGN_PARAGRAPH.CENTER)

    output = TEMPLATES_DIR / "procuracao_ad_judicia.docx"
    doc.save(str(output))
    print(f"  ✓ {output}")


def create_contrato() -> None:
    doc = Document()
    _set_margins(doc, top=3, bottom=2, left=3, right=2)

    # Título
    _add_paragraph(doc, "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, spacing_after=2)

    # Cláusula 1 - DAS PARTES
    _add_paragraph(doc, "CLÁUSULA 1ª — DAS PARTES", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "CONTRATANTE: {{ cliente_nome }}, inscrito(a) no CPF/CNPJ sob nº "
        "{{ cliente_cpf_cnpj }}, com endereço em {{ cliente_endereco }}."
    ))
    _add_paragraph(doc, (
        "CONTRATADO: {{ advogado_nome }}, inscrito(a) na OAB sob nº "
        "{{ advogado_oab }}, com escritório profissional na {{ advogado_endereco }}."
    ))

    # Cláusula 2 - DO OBJETO
    _add_paragraph(doc, "CLÁUSULA 2ª — DO OBJETO", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "O presente contrato tem por objeto a prestação de serviços advocatícios "
        "consistentes em: {{ objeto_contrato }}."
    ))

    # Cláusula 3 - DOS HONORÁRIOS
    _add_paragraph(doc, "CLÁUSULA 3ª — DOS HONORÁRIOS", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "Pelos serviços prestados, o CONTRATANTE pagará ao CONTRATADO a título "
        "de honorários advocatícios o valor de {{ honorarios }}."
    ))

    # Cláusula 4 - DAS DESPESAS
    _add_paragraph(doc, "CLÁUSULA 4ª — DAS DESPESAS", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "As despesas processuais, custas judiciais, taxas, emolumentos e "
        "quaisquer outros gastos necessários ao andamento do processo serão "
        "de responsabilidade exclusiva do CONTRATANTE, não estando incluídos "
        "nos honorários advocatícios."
    ))

    # Cláusula 5 - DA VIGÊNCIA
    _add_paragraph(doc, "CLÁUSULA 5ª — DA VIGÊNCIA", bold=True, spacing_after=0.5)
    _add_paragraph(doc, "O presente contrato vigorará {{ vigencia }}.")

    # Cláusula 6 - DA RESCISÃO
    _add_paragraph(doc, "CLÁUSULA 6ª — DA RESCISÃO", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "O presente contrato poderá ser rescindido por qualquer das partes, "
        "mediante notificação escrita com antecedência mínima de 30 (trinta) dias, "
        "sendo devidos os honorários proporcionais aos serviços já prestados."
    ))

    # Cláusula 7 - DO SIGILO
    _add_paragraph(doc, "CLÁUSULA 7ª — DO SIGILO", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "O CONTRATADO compromete-se a manter absoluto sigilo sobre todas "
        "as informações recebidas em razão do presente contrato, em conformidade "
        "com o dever de sigilo profissional previsto no Estatuto da Advocacia "
        "(Lei 8.906/1994) e no Código de Ética da OAB."
    ))

    # Cláusula 8 - DO FORO
    _add_paragraph(doc, "CLÁUSULA 8ª — DO FORO", bold=True, spacing_after=0.5)
    _add_paragraph(doc, (
        "Para dirimir quaisquer controvérsias oriundas deste contrato, "
        "as partes elegem o foro da Comarca de {{ foro }}."
    ))

    _add_paragraph(doc, "", spacing_after=1)

    # Local e data
    _add_paragraph(doc, (
        "E, por estarem justos e contratados, firmam o presente instrumento "
        "em 2 (duas) vias de igual teor e forma, na presença de 2 (duas) testemunhas."
    ))

    _add_paragraph(doc, "", spacing_after=1)
    _add_paragraph(doc, "{{ cidade }}, {{ data_extenso }}.", align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_paragraph(doc, "", spacing_after=2)

    # Assinaturas
    _add_paragraph(doc, "_" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
    _add_paragraph(doc, "CONTRATANTE: {{ cliente_nome }}", align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_paragraph(doc, "", spacing_after=1)

    _add_paragraph(doc, "_" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
    _add_paragraph(doc, "CONTRATADO: {{ advogado_nome }}", align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_paragraph(doc, "", spacing_after=2)

    # Testemunhas
    _add_paragraph(doc, "Testemunhas:", bold=True)
    _add_paragraph(doc, "", spacing_after=1)
    _add_paragraph(doc, "_" * 40 + "          " + "_" * 40, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
    _add_paragraph(doc, "Nome:                                                    Nome:", spacing_after=0)
    _add_paragraph(doc, "CPF:                                                       CPF:")

    output = TEMPLATES_DIR / "contrato_honorarios.docx"
    doc.save(str(output))
    print(f"  ✓ {output}")


if __name__ == "__main__":
    print("Gerando templates DOCX...")
    create_procuracao()
    create_contrato()
    print("Templates gerados com sucesso!")
